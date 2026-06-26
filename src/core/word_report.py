from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

from PIL import Image

from core.execution_log import desktop_scoped_path, safe_filename, safe_folder_name
from core.ocr_validation import (
    validate_applist_evidence_with_windows_ocr,
    validate_7zip_programs_and_features_with_windows_ocr,
    validate_adobe_acrobat_programs_and_features_with_windows_ocr,
    validate_edge_browser_version_with_windows_ocr,
    validate_fslogix_apps_programs_and_features_with_windows_ocr,
    validate_microsoft_office_programs_and_features_with_windows_ocr,
    validate_openjdk_jre_programs_and_features_with_windows_ocr,
    validate_microsoft_project_programs_and_features_with_windows_ocr,
    validate_microsoft_visio_programs_and_features_with_windows_ocr,
    validate_policy_pac_with_windows_ocr,
    validate_webview_version_with_windows_ocr,
)
from core.test_categories import (
    IAT_EVIDENCE_FOLDER,
    MANDATORY_EVIDENCE_FOLDER,
    POST_COMPLETE_ZSCALER_TEST_NAME,
    SHAKEDOWN_EVIDENCE_FOLDER,
    SILO43_EVIDENCE_FOLDER,
    is_silo43_desktop,
    is_success_status,
)


REPORT_STRUCTURE = [
    (
        "Mandatory Testcases",
        "mandatory",
        MANDATORY_EVIDENCE_FOLDER,
        [
            ("Hostname & IP address", ["Hostname_and_IP_Evidence"]),
            ("Edge & Edge WebView Versions", ["edge_evidence", "webview_evidence"]),
            ("ZScaler Services", ["zscaler_evidence", "zscaler_services_2"]),
            ("Office Applications Launch", ["powerpnt_evidence"]),
            ("Google and Yahoo Web Access", ["google_evidence", "yahoo_evidence"]),
            ("APP List (if applicable)", ["applist_evidence"]),
        ],
    ),
    (
        "SD: Shakedown Testcases",
        "shakedown",
        SHAKEDOWN_EVIDENCE_FOLDER,
        [
            ("Desktop Availability", ["desktop_availability"]),
            ("OneDrive Sync", ["onedrive_sync"]),
            ("Edge Sync", ["edge_sync", "edge_browser_version"]),
            ("Proxy PAC", ["policy_pac_1", "policy_pac_2"]),
            ("Windows Version", ["winver"]),
            ("Local and Network Drives", ["local_network_drives", "local_network_drives_deleted"]),
            ("FSLogix Profile Log Check", ["fslogix_profile_log"]),
            ("Verification of Unnecessary Files in C:\\Temp", ["temp_files"]),
        ],
    ),
    (
        "IAT: Core Applications Test",
        "iat",
        IAT_EVIDENCE_FOLDER,
        [
            ("7-Zip", ["7-zip_evidence"]),
            ("Adobe Acrobat Reader", ["adobe_acrobat_evidence"]),
            ("Microsoft Office", ["Microsoft_Office_evidence"]),
            ("Microsoft Visio", ["Microsoft_Visio_evidence"]),
            ("Microsoft Project", ["Microsoft_Project_evidence"]),
            ("Citrix VDA", ["citrix_vda_evidence"]),
            ("OpenJDK / JRE", ["OpenJDK_JRE_evidence"]),
            ("FSLogix", ["fslogix_apps_evidence"]),
        ],
    ),
    (
        "Silo 43 Testcases",
        "silo43",
        SILO43_EVIDENCE_FOLDER,
        [
            ("Oracle 12 bin path", ["silo43_oracle_12_bin_path_evidence"]),
            ("Nice Env Variables", ["silo43_oracle_12_bin_path_evidence"]),
            ("C:\\apps\\vls", ["silo43_vls_privilege_warning_evidence"]),
            ("Ping prod.dvfs.com", ["silo43_ping_prod_dvfs_evidence"]),
            ("C:\\BAD folder", ["silo43_bad_folder_evidence"]),
        ],
    ),
]

REPORT_TESTCASE_PREFIXES = {
    "Hostname_and_IP_Evidence": ["Hostname_and_IP_Evidence"],
    "Edge_WebView_Version_Evidence": ["webview_evidence"],
    "Edge_Browser_Version_Evidence": ["edge_evidence"],
    "Zscaler_Services_Evidence": ["zscaler_evidence"],
    POST_COMPLETE_ZSCALER_TEST_NAME: ["zscaler_services_2"],
    "Google_and_Yahoo_Web_Access_Evidence": ["google_evidence", "yahoo_evidence"],
    "Office_Applications_Launch": ["powerpnt_evidence"],
    "Applist_Validation_Evidence": ["applist_evidence"],
    "Shakedown_Desktop_Availability_Evidence": ["desktop_availability"],
    "Shakedown_OneDrive_Sync_Evidence": ["onedrive_sync"],
    "Shakedown_Edge_Sync_Evidence": ["edge_sync", "edge_browser_version"],
    "Shakedown_Edge_Policy_PAC_Evidence": ["policy_pac_1", "policy_pac_2"],
    "Shakedown_Windows_Version_Evidence": ["winver"],
    "Shakedown_Local_Network_Drives_Evidence": ["local_network_drives", "local_network_drives_deleted"],
    "Shakedown_FSLogix_Profile_Log_Evidence": ["fslogix_profile_log"],
    "Shakedown_Temp_Folder_Evidence": ["temp_files"],
    "IAT_Core_Application_Test_Evidence": [
        "7-zip_evidence",
        "adobe_acrobat_evidence",
        "Microsoft_Office_evidence",
        "Microsoft_Visio_evidence",
        "Microsoft_Project_evidence",
        "citrix_vda_evidence",
        "OpenJDK_JRE_evidence",
        "fslogix_apps_evidence",
    ],
    "Silo43_Oracle_12_Bin_Path_Evidence": ["silo43_oracle_12_bin_path_evidence"],
    "Silo43_Nice_Env_Variables_Evidence": ["silo43_oracle_12_bin_path_evidence"],
    "Silo43_VLS_Privilege_Warning_Evidence": ["silo43_vls_privilege_warning_evidence"],
    "Silo43_Ping_Prod_DVFS_Evidence": ["silo43_ping_prod_dvfs_evidence"],
    "Silo43_BAD_Folder_Evidence": ["silo43_bad_folder_evidence"],
}

APPLIST_SECTION_TITLE = "APP List (if applicable)"
APPLIST_NOT_OK_SECTION_TITLE = 'APP List (If applicable): Few apps are "NOT OK".'
POLICY_PAC_SECTION_TITLE = "Proxy PAC"
POLICY_PAC_NOT_OK_SECTION_TITLE = 'Policy PAC: Showing "NOT OK".'
SEVEN_ZIP_SECTION_TITLE = "7-Zip"
SEVEN_ZIP_NOT_AVAILABLE_TITLE = "7-Zip: Not available."
ADOBE_ACROBAT_SECTION_TITLE = "Adobe Acrobat Reader"
ADOBE_ACROBAT_NOT_AVAILABLE_TITLE = "Adobe Acrobat Reader: Not available."
MICROSOFT_OFFICE_SECTION_TITLE = "Microsoft Office"
MICROSOFT_OFFICE_NOT_AVAILABLE_TITLE = "Microsoft Office: Not available."
MICROSOFT_VISIO_SECTION_TITLE = "Microsoft Visio"
MICROSOFT_VISIO_NOT_AVAILABLE_TITLE = "Microsoft Visio: Not available."
MICROSOFT_PROJECT_SECTION_TITLE = "Microsoft Project"
MICROSOFT_PROJECT_NOT_AVAILABLE_TITLE = "Microsoft Project: Not available."
OPENJDK_JRE_SECTION_TITLE = "OpenJDK / JRE"
OPENJDK_JRE_NOT_AVAILABLE_TITLE = "OpenJDK / JRE: Not available."
FSLOGIX_SECTION_TITLE = "FSLogix"
FSLOGIX_NOT_AVAILABLE_TITLE = "FSLogix: Not available."


def generate_complete_testing_report(log_path: Path, screenshots_base_dir: Path, desktop_name: str) -> Path:
    payload = _read_payload(log_path)
    screenshots_root = desktop_scoped_path(screenshots_base_dir, desktop_name)
    return _generate_report_from_payload_and_folders(
        payload=payload,
        screenshots_root=screenshots_root,
        desktop_name=desktop_name,
        include_missing_subsections=True,
    )


def generate_report_from_screenshots(screenshots_base_dir: Path, desktop_name: str) -> Path:
    screenshots_root = desktop_scoped_path(screenshots_base_dir, desktop_name)
    return _generate_report_from_payload_and_folders(
        payload={},
        screenshots_root=screenshots_root,
        desktop_name=desktop_name,
        include_missing_subsections=False,
    )


def _generate_report_from_payload_and_folders(
    payload: dict,
    screenshots_root: Path,
    desktop_name: str,
    include_missing_subsections: bool,
) -> Path:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    evidence_root = screenshots_root.parent
    evidence_root.mkdir(parents=True, exist_ok=True)
    report_path = evidence_root / f"{safe_folder_name(desktop_name)}_Testing_.docx"

    doc = Document()
    _configure_document(doc, Pt, Inches)
    max_width_inches = _content_width_inches(doc)

    title = doc.add_paragraph()
    title_format = title.paragraph_format
    title_format.space_after = Pt(12)
    title_run = title.add_run(f"{desktop_name} Testing:")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(14)

    report_structure = _report_structure_for_desktop(desktop_name)
    report_payload = _payload_with_latest_evidence_status(payload, screenshots_root, desktop_name)

    summary_entries = _execution_summary_entries(report_payload, screenshots_root, desktop_name)
    if summary_entries:
        _add_execution_summary(doc, summary_entries, Pt, WD_TABLE_ALIGNMENT)

    failures = _failure_entries(report_payload, desktop_name)
    if failures:
        _add_failure_summary(doc, failures, Pt, WD_TABLE_ALIGNMENT)

    for section_title, payload_key, _folder_name, subsections in report_structure:
        phase_paths = _phase_screenshots(report_payload, payload_key)
        folder_paths = _folder_screenshots(screenshots_root / _folder_name)
        rendered_subsections: list[tuple[str, str, list[Path]]] = []
        for subsection_title, prefixes in subsections:
            images = _matching_images([*phase_paths, *folder_paths], prefixes)
            if not images and subsection_title == APPLIST_SECTION_TITLE:
                continue
            report_title = _report_subsection_title(report_payload, subsection_title, images)
            if images or include_missing_subsections:
                rendered_subsections.append((subsection_title, report_title, images))

        if not rendered_subsections:
            continue

        _add_section_heading(doc, section_title, Pt)
        for subsection_title, report_title, images in rendered_subsections:
            _add_subsection_heading(doc, report_title, Pt)
            if images:
                if subsection_title == "Edge & Edge WebView Versions":
                    _add_edge_version_summary(doc, images, Pt, report_payload)
                for image_path in images:
                    _add_image(doc, image_path, max_width_inches, Inches, WD_ALIGN_PARAGRAPH)
            else:
                _add_missing_note(doc, Pt)

    if not any(
        _folder_screenshots(screenshots_root / folder_name)
        for _section_title, _payload_key, folder_name, _subsections in report_structure
    ):
        _add_missing_note(doc, Pt)

    doc.save(report_path)
    return report_path


def _read_payload(log_path: Path) -> dict:
    with log_path.open("r", encoding="utf-8") as file:
        return json.load(file)


def _configure_document(doc, Pt, Inches) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)


def _content_width_inches(doc) -> float:
    section = doc.sections[0]
    content_width_emu = section.page_width - section.left_margin - section.right_margin
    return content_width_emu / 914400


def _add_section_heading(doc, text: str, Pt) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(13)


def _add_subsection_heading(doc, text: str, Pt) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)


def _add_missing_note(doc, Pt) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("No screenshot captured in this execution.")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)


def _add_edge_version_summary(doc, images: list[Path], Pt, payload: dict | None = None) -> None:
    edge_version = _extract_edge_version_from_payload(payload, "edge_browser") or _extract_report_version(
        images,
        "edge_evidence",
    )
    webview_version = _extract_edge_version_from_payload(payload, "webview") or _extract_report_version(
        images,
        "webview_evidence",
    )
    for label, version in (
        ("Edge", edge_version),
        ("Edge Webview", webview_version),
    ):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        label_run.bold = True
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(10)
        value_run = paragraph.add_run(version or "Not detected")
        value_run.font.name = "Calibri"
        value_run.font.size = Pt(10)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)


def _add_execution_summary(doc, entries: list[tuple[str, str]], Pt, WD_TABLE_ALIGNMENT) -> None:
    _add_section_heading(doc, "Execution Summary", Pt)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Metric", "Value"]
    for cell, header in zip(table.rows[0].cells, headers):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    for label, value in entries:
        row = table.add_row().cells
        for cell, text in zip(row, (label, value)):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
    doc.add_paragraph()


def _add_failure_summary(doc, failures: list[dict[str, str]], Pt, WD_TABLE_ALIGNMENT) -> None:
    _add_section_heading(doc, "Failure Summary", Pt)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Section", "Testcase", "Status", "Error / Log Path", "Screenshots"]
    for cell, header in zip(table.rows[0].cells, headers):
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(header)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)

    for failure in failures:
        row = table.add_row().cells
        values = [
            failure["section"],
            failure["test_case"],
            failure["status"],
            failure["detail"],
            failure["screenshots"],
        ]
        for cell, value in zip(row, values):
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(value)
            run.font.name = "Calibri"
            run.font.size = Pt(9)
    doc.add_paragraph()


def _add_image(doc, image_path: Path, max_width_inches: float, Inches, WD_ALIGN_PARAGRAPH) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Inches(0.12)
    width_inches = _fit_image_width_inches(image_path, max_width_inches)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    doc.add_paragraph()


def _fit_image_width_inches(image_path: Path, max_width_inches: float) -> float:
    try:
        with Image.open(image_path) as image:
            dpi = image.info.get("dpi", (96, 96))[0] or 96
            native_width_inches = image.width / dpi
    except OSError:
        native_width_inches = max_width_inches
    return min(native_width_inches, max_width_inches)


def _phase_screenshots(payload: dict, phase_key: str) -> list[Path]:
    phase = payload.get(phase_key, {})
    results = phase.get("individual_results", [])
    paths: list[Path] = []
    for result in results:
        for value in result.get("screenshots", []):
            path = Path(value)
            if path.exists():
                paths.append(path)
    return paths


def _folder_screenshots(folder: Path) -> list[Path]:
    if not folder.exists():
        return []
    return sorted(
        (path for path in folder.glob("*.png") if path.is_file()),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )


def _report_structure_for_desktop(desktop_name: str) -> list[tuple[str, str, str, list[tuple[str, list[str]]]]]:
    if is_silo43_desktop(desktop_name):
        return list(REPORT_STRUCTURE)
    return [section for section in REPORT_STRUCTURE if section[1] != "silo43"]


def _phase_labels_for_desktop(desktop_name: str) -> list[tuple[str, str]]:
    phase_labels = [
        ("mandatory", "Mandatory"),
        ("shakedown", "Shakedown"),
        ("iat", "IAT"),
        ("silo43", "Silo 43"),
    ]
    if not is_silo43_desktop(desktop_name):
        phase_labels = [item for item in phase_labels if item[0] != "silo43"]
    return phase_labels


def _payload_with_latest_evidence_status(payload: dict, screenshots_root: Path, desktop_name: str) -> dict:
    if not payload:
        return {}

    updated = dict(payload)
    phase_statuses: dict[str, str] = {}

    for phase_key, _label in _phase_labels_for_desktop(desktop_name):
        phase = payload.get(phase_key, {})
        if not isinstance(phase, dict):
            continue

        new_phase = dict(phase)
        new_results = []
        result_statuses: list[str] = []
        for result in phase.get("individual_results", []):
            if not isinstance(result, dict):
                new_results.append(result)
                continue

            new_result = dict(result)
            latest_status = _latest_evidence_status_for_result(new_result, screenshots_root)
            if latest_status:
                new_result["status"] = latest_status
                if is_success_status(latest_status):
                    new_result["error"] = None
            result_statuses.append(str(new_result.get("status") or "Unknown"))
            new_results.append(new_result)

        new_phase["individual_results"] = new_results
        new_phase_status = _phase_status_from_result_statuses(result_statuses, str(phase.get("status") or ""))
        if new_phase_status:
            new_phase["status"] = new_phase_status
            phase_statuses[phase_key] = new_phase_status
        updated[phase_key] = new_phase

    overall_status = _overall_status_from_phase_statuses(
        list(phase_statuses.values()),
        str(payload.get("overall_execution_result") or payload.get("status") or ""),
    )
    if overall_status:
        if "overall_execution_result" in updated or payload.get("overall_execution_result"):
            updated["overall_execution_result"] = overall_status
        elif "status" in updated:
            updated["status"] = overall_status
        else:
            updated["overall_execution_result"] = overall_status

    return updated


def _latest_evidence_status_for_result(result: dict, screenshots_root: Path) -> str | None:
    if not screenshots_root.exists():
        return None

    prefixes = _evidence_prefixes_for_result(result)
    if not prefixes:
        return None

    latest_statuses = []
    for prefix in prefixes:
        latest_image = _latest_screenshot_for_prefix(screenshots_root, prefix)
        if latest_image is None:
            return None
        status = _status_from_screenshot_filename(latest_image.name)
        if not status:
            return None
        latest_statuses.append(status)

    if any(status == "Fail" for status in latest_statuses):
        return "Fail"
    if latest_statuses and all(status == "Pass" for status in latest_statuses):
        return "Pass"
    return None


def _evidence_prefixes_for_result(result: dict) -> list[str]:
    test_case = str(result.get("test_case") or "")
    prefixes = REPORT_TESTCASE_PREFIXES.get(test_case)
    if prefixes:
        return list(prefixes)

    extracted: list[str] = []
    for key in ("screenshots", "evidence_screenshots"):
        raw_value = result.get(key)
        if isinstance(raw_value, list):
            for value in raw_value:
                prefix = _prefix_from_screenshot_name(Path(str(value)).name)
                if prefix:
                    extracted.append(prefix)
    for key in ("screenshot", "manual_confirmation_screenshot"):
        raw_value = result.get(key)
        if raw_value:
            prefix = _prefix_from_screenshot_name(Path(str(raw_value)).name)
            if prefix:
                extracted.append(prefix)
    return list(dict.fromkeys(extracted))


def _latest_screenshot_for_prefix(screenshots_root: Path, prefix: str) -> Path | None:
    prefix_key = prefix.casefold()
    matches = [
        path
        for path in screenshots_root.rglob("*.png")
        if path.is_file() and (_prefix_from_screenshot_name(path.name) or "").casefold() == prefix_key
    ]
    if not matches:
        return None
    return max(matches, key=lambda path: path.stat().st_mtime)


def _prefix_from_screenshot_name(filename: str) -> str | None:
    match = re.match(r"(.+?)_(?:Pass|Fail)_\d{8}_\d{6}\.png$", filename, flags=re.IGNORECASE)
    return match.group(1) if match else None


def _status_from_screenshot_filename(filename: str) -> str | None:
    name = filename.casefold()
    if "_fail_" in name:
        return "Fail"
    if "_pass_" in name:
        return "Pass"
    return None


def _phase_status_from_result_statuses(statuses: list[str], fallback: str) -> str:
    clean_statuses = [status for status in statuses if status]
    if not clean_statuses:
        return fallback
    if any(status == "Fail" for status in clean_statuses):
        return "Fail"
    if all(status == "Skipped" for status in clean_statuses):
        return "Skipped"
    if all(is_success_status(status) for status in clean_statuses):
        return "Pass"
    return fallback


def _overall_status_from_phase_statuses(statuses: list[str], fallback: str) -> str:
    clean_statuses = [status for status in statuses if status]
    if not clean_statuses:
        return fallback
    if any(status == "Fail" for status in clean_statuses):
        return "Fail"
    if all(status == "Skipped" for status in clean_statuses):
        return "Skipped"
    if all(is_success_status(status) for status in clean_statuses):
        return "Pass"
    return fallback


def _execution_summary_entries(payload: dict, screenshots_root: Path, desktop_name: str) -> list[tuple[str, str]]:
    phase_labels = _phase_labels_for_desktop(desktop_name)
    result_statuses: list[str] = []
    entries: list[tuple[str, str]] = []

    overall_status = str(payload.get("overall_execution_result") or payload.get("status") or "").strip()
    if overall_status:
        entries.append(("Overall Status", overall_status))

    for phase_key, label in phase_labels:
        phase = payload.get(phase_key, {}) if isinstance(payload.get(phase_key), dict) else {}
        status = str(phase.get("status") or "").strip()
        if status:
            entries.append((f"{label} Status", status))
        for result in phase.get("individual_results", []):
            if isinstance(result, dict):
                result_statuses.append(str(result.get("status") or "Unknown"))

    if result_statuses:
        entries.extend(
            [
                ("Passed Testcases", str(sum(1 for status in result_statuses if is_success_status(status)))),
                ("Failed Testcases", str(sum(1 for status in result_statuses if status == "Fail"))),
                ("Skipped Testcases", str(sum(1 for status in result_statuses if status == "Skipped"))),
                ("Unknown Testcases", str(sum(1 for status in result_statuses if status in {"", "Unknown"}))),
            ]
        )

    screenshots = list(screenshots_root.rglob("*.png")) if screenshots_root.exists() else []
    if screenshots:
        entries.append(("Screenshots Indexed", str(len(screenshots))))
        entries.append(("Failed Screenshots", str(sum(1 for path in screenshots if "_Fail_" in path.name))))

    return entries


def _failure_entries(payload: dict, desktop_name: str) -> list[dict[str, str]]:
    entries: list[dict[str, str]] = []
    labels = {
        "mandatory": "Mandatory",
        "shakedown": "Shakedown",
        "iat": "IAT",
        "silo43": "Silo 43",
    }
    if not is_silo43_desktop(desktop_name):
        labels.pop("silo43", None)
    for phase_key, label in labels.items():
        phase = payload.get(phase_key, {})
        for result in phase.get("individual_results", []):
            status = result.get("status", "Unknown")
            if is_success_status(status):
                continue
            screenshots = [str(value) for value in result.get("screenshots", [])]
            detail_parts = []
            if result.get("error"):
                detail_parts.append(str(result["error"]))
            if result.get("log_path"):
                detail_parts.append(str(result["log_path"]))
            entries.append(
                {
                    "section": label,
                    "test_case": str(result.get("test_case", "Unknown")),
                    "status": str(status),
                    "detail": "\n".join(detail_parts) if detail_parts else "-",
                    "screenshots": "\n".join(screenshots) if screenshots else "-",
                }
            )
    return entries


def _report_subsection_title(payload: dict, subsection_title: str, images: list[Path]) -> str:
    if subsection_title == SEVEN_ZIP_SECTION_TITLE:
        return _seven_zip_report_title(payload, images)
    if subsection_title == ADOBE_ACROBAT_SECTION_TITLE:
        return _adobe_acrobat_report_title(payload, images)
    if subsection_title == MICROSOFT_OFFICE_SECTION_TITLE:
        return _microsoft_office_report_title(payload, images)
    if subsection_title == MICROSOFT_VISIO_SECTION_TITLE:
        return _microsoft_visio_report_title(payload, images)
    if subsection_title == MICROSOFT_PROJECT_SECTION_TITLE:
        return _microsoft_project_report_title(payload, images)
    if subsection_title == OPENJDK_JRE_SECTION_TITLE:
        return _openjdk_jre_report_title(payload, images)
    if subsection_title == FSLOGIX_SECTION_TITLE:
        return _fslogix_report_title(payload, images)
    if subsection_title == APPLIST_SECTION_TITLE and _applist_not_ok_found(payload, images):
        return APPLIST_NOT_OK_SECTION_TITLE
    if subsection_title == POLICY_PAC_SECTION_TITLE and _policy_pac_not_ok_found(payload, images):
        return POLICY_PAC_NOT_OK_SECTION_TITLE
    return subsection_title


def _seven_zip_report_title(payload: dict, images: list[Path]) -> str:
    seven_zip_image = next((path for path in images if _matches_evidence_prefix(path, "7-zip_evidence")), None)
    if seven_zip_image is not None:
        try:
            result = validate_7zip_programs_and_features_with_windows_ocr(seven_zip_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return f"7-Zip: {result.version}" if result.version else SEVEN_ZIP_NOT_AVAILABLE_TITLE

    metadata = _seven_zip_metadata(payload)
    if metadata.get("seven_zip_available") is True:
        version = str(metadata.get("seven_zip_version", "")).strip()
        if version:
            return f"7-Zip: {version}"
    if metadata.get("seven_zip_search_term"):
        return SEVEN_ZIP_NOT_AVAILABLE_TITLE
    return SEVEN_ZIP_SECTION_TITLE


def _adobe_acrobat_report_title(payload: dict, images: list[Path]) -> str:
    adobe_image = next(
        (path for path in images if _matches_evidence_prefix(path, "adobe_acrobat_evidence")),
        None,
    )
    if adobe_image is not None:
        try:
            result = validate_adobe_acrobat_programs_and_features_with_windows_ocr(adobe_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return (
                f"Adobe Acrobat Reader: {result.version}"
                if result.version
                else ADOBE_ACROBAT_NOT_AVAILABLE_TITLE
            )

    metadata = _adobe_acrobat_metadata(payload)
    if metadata.get("adobe_acrobat_available") is True:
        version = str(metadata.get("adobe_acrobat_version", "")).strip()
        if version:
            return f"Adobe Acrobat Reader: {version}"
    if metadata.get("adobe_acrobat_search_term"):
        return ADOBE_ACROBAT_NOT_AVAILABLE_TITLE
    return ADOBE_ACROBAT_SECTION_TITLE


def _microsoft_office_report_title(payload: dict, images: list[Path]) -> str:
    office_image = next(
        (path for path in images if _matches_evidence_prefix(path, "Microsoft_Office_evidence")),
        None,
    )
    if office_image is not None:
        try:
            result = validate_microsoft_office_programs_and_features_with_windows_ocr(office_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return (
                f"Microsoft Office: {result.version}"
                if result.version
                else MICROSOFT_OFFICE_NOT_AVAILABLE_TITLE
            )

    metadata = _iat_metadata(payload)
    if metadata.get("microsoft_office_available") is True:
        version = str(metadata.get("microsoft_office_version", "")).strip()
        if version:
            return f"Microsoft Office: {version}"
    if metadata.get("microsoft_office_search_term"):
        return MICROSOFT_OFFICE_NOT_AVAILABLE_TITLE
    return MICROSOFT_OFFICE_SECTION_TITLE


def _microsoft_visio_report_title(payload: dict, images: list[Path]) -> str:
    visio_image = next(
        (path for path in images if _matches_evidence_prefix(path, "Microsoft_Visio_evidence")),
        None,
    )
    if visio_image is not None:
        try:
            result = validate_microsoft_visio_programs_and_features_with_windows_ocr(visio_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return (
                f"Microsoft Visio: {result.version}"
                if result.version
                else MICROSOFT_VISIO_NOT_AVAILABLE_TITLE
            )

    metadata = _iat_metadata(payload)
    if metadata.get("microsoft_visio_available") is True:
        version = str(metadata.get("microsoft_visio_version", "")).strip()
        if version:
            return f"Microsoft Visio: {version}"
    if metadata.get("microsoft_visio_search_term"):
        return MICROSOFT_VISIO_NOT_AVAILABLE_TITLE
    return MICROSOFT_VISIO_SECTION_TITLE


def _microsoft_project_report_title(payload: dict, images: list[Path]) -> str:
    project_image = next(
        (path for path in images if _matches_evidence_prefix(path, "Microsoft_Project_evidence")),
        None,
    )
    if project_image is not None:
        try:
            result = validate_microsoft_project_programs_and_features_with_windows_ocr(project_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return (
                f"Microsoft Project: {result.version}"
                if result.version
                else MICROSOFT_PROJECT_NOT_AVAILABLE_TITLE
            )

    metadata = _iat_metadata(payload)
    if metadata.get("microsoft_project_available") is True:
        version = str(metadata.get("microsoft_project_version", "")).strip()
        if version:
            return f"Microsoft Project: {version}"
    if metadata.get("microsoft_project_search_term"):
        return MICROSOFT_PROJECT_NOT_AVAILABLE_TITLE
    return MICROSOFT_PROJECT_SECTION_TITLE


def _openjdk_jre_report_title(payload: dict, images: list[Path]) -> str:
    jre_image = next(
        (path for path in images if _matches_evidence_prefix(path, "OpenJDK_JRE_evidence")),
        None,
    )
    if jre_image is not None:
        try:
            result = validate_openjdk_jre_programs_and_features_with_windows_ocr(jre_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return f"OpenJDK / JRE: {result.version}" if result.version else OPENJDK_JRE_NOT_AVAILABLE_TITLE

    metadata = _iat_metadata(payload)
    if metadata.get("openjdk_jre_available") is True:
        version = str(metadata.get("openjdk_jre_version", "")).strip()
        if version:
            return f"OpenJDK / JRE: {version}"
    if metadata.get("openjdk_jre_search_term"):
        return OPENJDK_JRE_NOT_AVAILABLE_TITLE
    return OPENJDK_JRE_SECTION_TITLE


def _fslogix_report_title(payload: dict, images: list[Path]) -> str:
    fslogix_image = next(
        (path for path in images if _matches_evidence_prefix(path, "fslogix_apps_evidence")),
        None,
    )
    if fslogix_image is not None:
        try:
            result = validate_fslogix_apps_programs_and_features_with_windows_ocr(fslogix_image)
        except Exception:
            result = None
        if result is not None and result.valid:
            return f"FSLogix: {result.version}" if result.version else FSLOGIX_NOT_AVAILABLE_TITLE

    metadata = _iat_metadata(payload)
    if metadata.get("fslogix_apps_available") is True:
        version = str(metadata.get("fslogix_apps_version", "")).strip()
        if version:
            return f"FSLogix: {version}"
    if metadata.get("fslogix_apps_search_term"):
        return FSLOGIX_NOT_AVAILABLE_TITLE
    return FSLOGIX_SECTION_TITLE


def _applist_not_ok_found(payload: dict, images: list[Path]) -> bool:
    applist_image = next(
        (path for path in images if _matches_evidence_prefix(path, "applist_evidence")),
        None,
    )
    image_signal = None
    if applist_image is not None:
        try:
            result = validate_applist_evidence_with_windows_ocr(applist_image)
        except Exception:
            result = None
        if result is not None:
            image_signal = _applist_not_ok_signal_from_ocr_text(result.raw_text)

    if image_signal is False:
        return False
    metadata_signal = _applist_not_ok_signal_from_payload(payload)
    if metadata_signal is not None:
        return metadata_signal
    if image_signal is not None:
        return image_signal
    return False


def _applist_not_ok_signal_from_payload(payload: dict) -> bool | None:
    mandatory = payload.get("mandatory", {})
    for result in mandatory.get("individual_results", []):
        if result.get("test_case") != "Applist_Validation_Evidence":
            continue
        metadata = result.get("metadata") or {}
        if "applist_not_ok_found" in metadata:
            return metadata.get("applist_not_ok_found") is True
    return None


def _applist_not_ok_signal_from_ocr_text(raw_text: str) -> bool | None:
    if not raw_text.strip():
        return None
    compact_text = re.sub(r"[^a-z0-9]+", "", raw_text.casefold())
    not_ok_visible = "notok" in compact_text or "not0k" in compact_text
    if not not_ok_visible:
        return False
    not_ok_missing_dialog = (
        "cannotfindnotok" in compact_text
        or "cannotfindnot0k" in compact_text
        or "cantfindnotok" in compact_text
        or "couldnotfindnotok" in compact_text
    )
    if not_ok_missing_dialog:
        return False
    return True


def _policy_pac_not_ok_found(payload: dict, images: list[Path]) -> bool:
    policy_images = [path for path in images if _matches_evidence_prefix(path, "policy_pac_1") or _matches_evidence_prefix(path, "policy_pac_2")]
    for policy_image in policy_images:
        try:
            result = validate_policy_pac_with_windows_ocr(policy_image)
        except Exception:
            result = None
        if result is not None and not result.valid:
            reason = result.reason.casefold()
            if "not ok" in reason or "error" in reason:
                return True

    shakedown = payload.get("shakedown", {})
    for result in shakedown.get("individual_results", []):
        if result.get("test_case") != "Shakedown_Edge_Policy_PAC_Evidence":
            continue
        metadata = result.get("metadata") or {}
        if metadata.get("policy_pac_not_ok_found") is True or metadata.get("policy_pac_error_found") is True:
            return True
    return False


def _seven_zip_metadata(payload: dict) -> dict:
    return _iat_metadata(payload)


def _adobe_acrobat_metadata(payload: dict) -> dict:
    return _iat_metadata(payload)


def _iat_metadata(payload: dict) -> dict:
    iat = payload.get("iat", {})
    for result in iat.get("individual_results", []):
        if result.get("test_case") != "IAT_Core_Application_Test_Evidence":
            continue
        metadata = result.get("metadata") or {}
        if isinstance(metadata, dict):
            return metadata
    return {}


def _matching_images(paths: Iterable[Path], prefixes: list[str]) -> list[Path]:
    selected: list[Path] = []
    unique_paths = []
    seen = set()
    for path in paths:
        resolved = str(path)
        if resolved in seen:
            continue
        seen.add(resolved)
        unique_paths.append(path)

    for prefix in prefixes:
        matches = sorted(
            (path for path in unique_paths if _matches_evidence_prefix(path, prefix)),
            key=lambda path: path.stat().st_mtime,
            reverse=True,
        )
        if matches:
            selected.append(matches[0])
    return selected


def _matches_evidence_prefix(path: Path, prefix: str) -> bool:
    safe_prefix = safe_filename(prefix).casefold()
    pattern = re.compile(rf"^{re.escape(safe_prefix)}_(Pass|Fail)_[0-9]{{8}}_[0-9]{{6}}\.png$", re.IGNORECASE)
    return bool(pattern.match(path.name))


def _extract_report_version(images: list[Path], prefix: str) -> str:
    for image_path in images:
        if not _matches_evidence_prefix(image_path, prefix):
            continue
        try:
            if prefix == "edge_evidence":
                result = validate_edge_browser_version_with_windows_ocr(image_path)
            else:
                result = validate_webview_version_with_windows_ocr(image_path)
        except Exception:
            return ""
        if result.valid and result.version:
            return result.version
        return _extract_loose_registry_version(result.raw_text, prefix)
    return ""


_REPORT_VERSION_RE = re.compile(
    r"([0-9oOeE]{2,3}\s*\.\s*[0-9oOeE]{1,5}\s*\.\s*[0-9oOeE]{1,5}\s*\.\s*[0-9oOeE]{1,5})"
)


def _extract_edge_version_from_payload(payload: dict | None, product: str) -> str:
    if not payload:
        return ""

    mandatory = payload.get("mandatory") or {}
    results = mandatory.get("individual_results") or []
    if product == "webview":
        target_tests = {"Edge_WebView_Version_Evidence"}
        metadata_keys = ("webview_version", "edge_webview_version")
    else:
        target_tests = {"Edge_Browser_Version_Evidence", "Edge_WebView_Version_Evidence"}
        metadata_keys = ("edge_browser_version", "edge_version")

    for result in results:
        if not isinstance(result, dict):
            continue
        metadata = result.get("metadata") or {}
        for key in metadata_keys:
            version = _normalize_report_version(str(metadata.get(key) or ""))
            if version:
                return version

        test_case = str(result.get("test_case") or "")
        if test_case not in target_tests:
            continue
        log_path = Path(str(result.get("log_path") or ""))
        version = _extract_edge_version_from_log(log_path, product)
        if version:
            return version

    return ""


def _extract_edge_version_from_log(log_path: Path, product: str) -> str:
    if not log_path.exists():
        return ""
    try:
        payload = json.loads(log_path.read_text(encoding="utf-8"))
    except Exception:
        return ""

    for message in _log_step_messages(payload):
        version = _extract_edge_version_from_message(message, product)
        if version:
            return version
    return ""


def _log_step_messages(payload: dict) -> list[str]:
    messages: list[str] = []
    for step in payload.get("steps") or []:
        if isinstance(step, dict):
            messages.append(str(step.get("message") or ""))
        else:
            messages.append(str(step))
    return messages


def _extract_edge_version_from_message(message: str, product: str) -> str:
    lower = message.casefold()
    if product == "webview":
        if "webview" not in lower:
            return ""
    else:
        if "edge browser" not in lower and "microsoft edge version" not in lower:
            return ""
        if "webview" in lower:
            return ""

    for pattern in (
        re.compile(r"\bversion\s*=\s*" + _REPORT_VERSION_RE.pattern, re.IGNORECASE),
        re.compile(r"\bversion\s+" + _REPORT_VERSION_RE.pattern, re.IGNORECASE),
    ):
        match = pattern.search(message)
        if match:
            version = _normalize_report_version(match.group(1))
            if version:
                return version
    return ""


def _extract_loose_registry_version(raw_text: str, prefix: str) -> str:
    if not raw_text:
        return ""

    compact = re.sub(r"\s+", " ", raw_text).casefold()
    if prefix == "edge_evidence" and "webview" in compact:
        return ""

    pv_match = re.search(
        r"\bpv\b\s*[:=]?\s*" + _REPORT_VERSION_RE.pattern,
        raw_text,
        flags=re.IGNORECASE,
    )
    if pv_match:
        version = _normalize_report_version(pv_match.group(1))
        if version:
            return version

    match = _REPORT_VERSION_RE.search(raw_text)
    return _normalize_report_version(match.group(1)) if match else ""


def _normalize_report_version(value: str) -> str:
    version = re.sub(r"\s+", "", value or "")
    version = version.replace("O", "0").replace("o", "0").replace("E", "0").replace("e", "0")
    return version if re.fullmatch(r"\d{2,3}(?:\.\d{1,5}){3}", version) else ""
